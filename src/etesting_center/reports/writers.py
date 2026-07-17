from __future__ import annotations

import json
from dataclasses import asdict
from datetime import datetime
from html import escape
from pathlib import Path

from etesting_center import APP_NAME, APP_VERSION
from etesting_center.engine.models import ScanReport


def report_to_dict(report: ScanReport) -> dict:
    return asdict(report)


def write_report(report: ScanReport, output: Path, fmt: str) -> Path:
    fmt = fmt.lower()
    output.parent.mkdir(parents=True, exist_ok=True)
    if fmt == "json":
        output.write_text(json.dumps(report_to_dict(report), indent=2, ensure_ascii=False), encoding="utf-8")
    elif fmt == "html":
        output.write_text(render_html(report), encoding="utf-8")
    elif fmt == "txt":
        output.write_text(render_text(report), encoding="utf-8")
    elif fmt == "docx":
        write_docx(report, output)
    else:
        raise ValueError(f"Unsupported report format: {fmt}")
    return output


def render_text(report: ScanReport) -> str:
    lines = [
        f"{APP_NAME} {APP_VERSION} scan report",
        f"Target: {report.target}",
        f"Started: {report.started_at}",
        f"Finished: {report.finished_at}",
        f"YARA enabled: {report.yara_enabled}",
        "",
        "Summary",
        f"  Scanned: {report.summary.scanned}",
        f"  Safe: {report.summary.safe}",
        f"  Suspicious: {report.summary.suspicious}",
        f"  Malicious: {report.summary.malicious}",
        f"  Errors: {report.summary.errors}",
        "",
        "Results",
    ]
    for result in report.results:
        lines.extend(
            [
                f"[{result.risk.upper()}] {result.path}",
                f"  Score: {result.score}",
                f"  Type: {result.file_type}",
                f"  Size: {result.size}",
                f"  SHA-256: {result.sha256}",
            ]
        )
        if result.error:
            lines.append(f"  Error: {result.error}")
        for finding in result.findings:
            lines.append(f"  - {finding.engine}:{finding.rule} ({finding.confidence}) {finding.description}")
        lines.append("")
    return "\n".join(lines)


def render_html(report: ScanReport) -> str:
    rows = []
    for result in report.results:
        findings = "<br>".join(
            escape(f"{item.engine}:{item.rule} ({item.confidence}) {item.description}") for item in result.findings
        ) or "-"
        rows.append(
            "<tr>"
            f"<td><span class='badge {escape(result.risk)}'>{escape(result.risk)}</span></td>"
            f"<td>{escape(result.path)}</td>"
            f"<td>{result.score}</td>"
            f"<td>{escape(result.file_type)}</td>"
            f"<td>{result.size}</td>"
            f"<td class='hash'>{escape(result.sha256)}</td>"
            f"<td>{findings}</td>"
            "</tr>"
        )
    return f"""<!doctype html>
<html lang="en">
<head>
<meta charset="utf-8">
<title>{escape(APP_NAME)} Report</title>
<style>
body {{ margin: 0; font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", sans-serif; background: #f5f6f8; color: #1d1d1f; }}
main {{ max-width: 1180px; margin: 0 auto; padding: 36px 24px; }}
h1 {{ margin: 0; font-size: 30px; font-weight: 680; letter-spacing: 0; }}
.meta {{ color: #62666d; margin-top: 8px; line-height: 1.5; }}
.summary {{ display: grid; grid-template-columns: repeat(5, minmax(120px, 1fr)); gap: 12px; margin: 26px 0; }}
.metric {{ background: rgba(255,255,255,.82); border: 1px solid #e3e5e8; border-radius: 8px; padding: 14px; }}
.metric strong {{ display: block; font-size: 24px; }}
.metric span {{ color: #62666d; font-size: 13px; }}
table {{ width: 100%; border-collapse: collapse; background: rgba(255,255,255,.88); border: 1px solid #e2e4e8; border-radius: 8px; overflow: hidden; }}
th, td {{ padding: 10px 12px; border-bottom: 1px solid #edf0f2; text-align: left; vertical-align: top; font-size: 13px; }}
th {{ color: #62666d; font-weight: 600; background: #fbfbfc; }}
.hash {{ font-family: "Cascadia Mono", Consolas, monospace; word-break: break-all; }}
.badge {{ display: inline-block; min-width: 72px; text-align: center; border-radius: 999px; padding: 4px 8px; font-weight: 650; }}
.safe {{ background: #e7f5ee; color: #12663d; }}
.suspicious {{ background: #fff1cf; color: #805400; }}
.malicious {{ background: #ffe0df; color: #9b1c16; }}
</style>
</head>
<body>
<main>
<h1>{escape(APP_NAME)} scan report</h1>
<div class="meta">Target: {escape(report.target)}<br>Started: {escape(report.started_at)}<br>Finished: {escape(report.finished_at)}<br>YARA enabled: {report.yara_enabled}</div>
<section class="summary">
<div class="metric"><strong>{report.summary.scanned}</strong><span>Scanned</span></div>
<div class="metric"><strong>{report.summary.safe}</strong><span>Safe</span></div>
<div class="metric"><strong>{report.summary.suspicious}</strong><span>Suspicious</span></div>
<div class="metric"><strong>{report.summary.malicious}</strong><span>Malicious</span></div>
<div class="metric"><strong>{report.summary.errors}</strong><span>Errors</span></div>
</section>
<table>
<thead><tr><th>Risk</th><th>Path</th><th>Score</th><th>Type</th><th>Size</th><th>SHA-256</th><th>Findings</th></tr></thead>
<tbody>{''.join(rows)}</tbody>
</table>
</main>
</body>
</html>
"""


# ------------------------------------------------------------------
# Word 报告生成
# ------------------------------------------------------------------

def write_docx(report: ScanReport, output: Path) -> Path:
    """
    生成企业级 Word (.docx) 扫描报告。

    报告结构：
    1. 封面区 — 标题、版本、生成时间
    2. 扫描概要 — 目标、时间、YARA 状态
    3. 扫描统计 — 已扫描/安全/可疑/恶意/错误 数量
    4. 威胁详情表 — 文件名/路径/风险等级/评分/引擎命中
    5. 每条威胁的详细发现依据
    6. 建议措施
    7. 签名栏
    """
    from docx import Document
    from docx.enum.section import WD_ORIENT
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.shared import Cm, Inches, Pt, RGBColor
    from docx.shared import Emu

    doc = Document()

    # -- 全局样式设置 --
    style = doc.styles["Normal"]
    style.font.name = "微软雅黑"
    style.font.size = Pt(10.5)
    style.paragraph_format.space_after = Pt(6)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    # 页边距
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    # ---- 封面标题 ----
    _add_docx_heading(doc, APP_NAME, level=0, font_size=28, color=RGBColor(0x1F, 0x6F, 0xEB))
    _add_docx_paragraph(doc, f"版本 {APP_VERSION}  ·  威胁扫描报告", alignment=WD_ALIGN_PARAGRAPH.CENTER, font_size=13, color=RGBColor(0x62, 0x66, 0x6D))
    _add_docx_paragraph(doc, f"生成时间：{_format_local_time()}", alignment=WD_ALIGN_PARAGRAPH.CENTER, font_size=11, color=RGBColor(0x8B, 0x8E, 0x94))
    doc.add_paragraph("")  # 空行

    # ---- 分隔线 ----
    _add_horizontal_line(doc)

    # ---- 扫描概要 ----
    _add_docx_heading(doc, "一、扫描概要", level=1)
    info_data = [
        ("扫描目标", report.target),
        ("开始时间", _format_iso(report.started_at)),
        ("结束时间", _format_iso(report.finished_at)),
        ("YARA 规则", "已启用" if report.yara_enabled else "未启用"),
        ("扫描状态", "已取消" if report.cancelled else "已完成"),
    ]
    _add_docx_info_table(doc, info_data)

    # ---- 扫描统计 ----
    _add_docx_heading(doc, "二、扫描统计", level=1)
    stats = report.summary
    stat_data = [
        ("已扫描文件数", str(stats.scanned)),
        ("安全", str(stats.safe)),
        ("可疑", str(stats.suspicious)),
        ("恶意", str(stats.malicious)),
        ("错误", str(stats.errors)),
    ]
    _add_docx_stat_table(doc, stat_data)

    # ---- 威胁详情 ----
    _add_docx_heading(doc, "三、威胁详情", level=1)

    threats = [r for r in report.results if r.risk in ("malicious", "suspicious")]
    safe_count = len([r for r in report.results if r.risk == "safe"])

    if not threats:
        _add_docx_paragraph(doc, f"未发现可疑或恶意文件。安全文件共 {safe_count} 个。", font_size=11)
    else:
        _add_docx_paragraph(doc, f"共检出 {len(threats)} 个威胁项（恶意 {stats.malicious} / 可疑 {stats.suspicious}），安全文件 {safe_count} 个。以下是威胁详情：", font_size=11)
        doc.add_paragraph("")

        # 威胁汇总表
        threat_table = doc.add_table(rows=1, cols=6)
        threat_table.style = "Table Grid"
        threat_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        _set_table_header(threat_table.rows[0], ["文件名称", "风险等级", "评分", "类型", "大小", "路径"])

        for r in threats:
            row = threat_table.add_row()
            risk_text = _risk_cn(r.risk)
            row.cells[0].text = Path(r.path).name
            row.cells[1].text = risk_text
            row.cells[2].text = str(r.score)
            row.cells[3].text = r.file_type
            row.cells[4].text = _format_size(r.size)
            row.cells[5].text = r.path
            # 风险等级着色
            _set_cell_shading(row.cells[1], _risk_color_hex(r.risk))
            for cell in row.cells:
                _set_cell_font(cell, size=Pt(9))

        # 列宽调整
        _set_table_col_widths(threat_table, [Cm(2.2), Cm(1.5), Cm(1.0), Cm(1.2), Cm(1.5), Cm(7.6)])

        doc.add_paragraph("")

        # 每条威胁的引擎命中详情
        _add_docx_heading(doc, "3.1 引擎命中详情", level=2)
        for idx, r in enumerate(threats, start=1):
            _add_docx_heading(doc, f"#{idx}  {Path(r.path).name}", level=3)
            detail_data = [
                ("完整路径", r.path),
                ("风险等级", f"{_risk_cn(r.risk)}（评分 {r.score}）"),
                ("文件类型", r.file_type),
                ("文件大小", _format_size(r.size)),
                ("MD5", r.md5),
                ("SHA-256", r.sha256),
            ]
            _add_docx_info_table(doc, detail_data)

            if r.error:
                _add_docx_paragraph(doc, f"错误：{r.error}", color=RGBColor(0xB4, 0x23, 0x18))

            if r.findings:
                _add_docx_paragraph(doc, "检测依据：", font_size=10, bold=True)
                for f in r.findings:
                    _add_docx_paragraph(
                        doc,
                        f"  [{_risk_cn(f.severity)}] {f.engine}:{f.rule}（置信度 {f.confidence}）— {f.description}",
                        font_size=9.5,
                    )
            else:
                _add_docx_paragraph(doc, "未命中本地指标。", font_size=9.5, color=RGBColor(0x8B, 0x8E, 0x94))

    # ---- 建议措施 ----
    _add_docx_heading(doc, "四、建议措施", level=1)
    recommendations = [
        "恶意文件：建议立即删除或隔离，并通过 VirusTotal 等平台二次确认。",
        "可疑文件：建议上传至 VirusTotal 进行多引擎交叉验证，确认安全后再放行。",
        "无签名可执行文件：建议核实来源，优先使用具有有效数字签名的版本。",
        "定期更新检测规则库，保持 YARA 规则与哈希特征库为最新版本。",
        "对系统关键区域（TEMP、Downloads、启动项）定期执行快速扫描。",
    ]
    for i, rec in enumerate(recommendations, start=1):
        _add_docx_paragraph(doc, f"{i}. {rec}", font_size=10.5)

    # ---- 签名栏 ----
    _add_horizontal_line(doc)
    _add_docx_paragraph(doc, "", font_size=6)
    _add_docx_paragraph(doc, f"本报告由 {APP_NAME} {APP_VERSION} 自动生成", alignment=WD_ALIGN_PARAGRAPH.CENTER, font_size=9, color=RGBColor(0x8B, 0x8E, 0x94))
    _add_docx_paragraph(doc, "只读检测工具  ·  不修改、不隔离、不删除任何文件", alignment=WD_ALIGN_PARAGRAPH.CENTER, font_size=9, color=RGBColor(0x8B, 0x8E, 0x94))
    _add_docx_paragraph(doc, "", font_size=6)

    sign_data = [
        ("分析人员", ""),
        ("审核人员", ""),
        ("日期", ""),
    ]
    _add_docx_info_table(doc, sign_data)

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output))
    return output


# ------------------------------------------------------------------
# Word 报告辅助函数
# ------------------------------------------------------------------

def _add_docx_heading(doc, text: str, level: int = 1, font_size: int | None = None, color=None):
    """添加标题段落。level=0 为大封面标题。"""
    if level == 0:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(font_size or 28)
        run.font.color.rgb = color or RGBColor(0x1D, 0x1D, 0x1F)
        run.font.name = "微软雅黑"
        run.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    else:
        heading = doc.add_heading(text, level=min(level, 3))
        for run in heading.runs:
            run.font.name = "微软雅黑"
            run.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")


def _add_docx_paragraph(doc, text: str, alignment=None, font_size=None, color=None, bold: bool = False):
    """添加普通段落。"""
    p = doc.add_paragraph()
    if alignment is not None:
        p.alignment = alignment
    run = p.add_run(text)
    run.font.name = "微软雅黑"
    run.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    if font_size:
        run.font.size = Pt(font_size)
    if color:
        run.font.color.rgb = color
    if bold:
        run.bold = True
    return p


def _add_docx_info_table(doc, rows: list[tuple[str, str]]):
    """添加键值对信息表（两列，左列加粗标签）。"""
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (key, val) in enumerate(rows):
        table.rows[i].cells[0].text = key
        table.rows[i].cells[1].text = val
        _set_cell_font(table.rows[i].cells[0], size=Pt(10), bold=True)
        _set_cell_font(table.rows[i].cells[1], size=Pt(10))
        _set_cell_shading(table.rows[i].cells[0], "F0F3F8")
    _set_table_col_widths(table, [Cm(3.0), Cm(12.6)])


def _add_docx_stat_table(doc, rows: list[tuple[str, str]]):
    """添加统计信息表。"""
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (key, val) in enumerate(rows):
        table.rows[i].cells[0].text = key
        table.rows[i].cells[1].text = val
        _set_cell_font(table.rows[i].cells[0], size=Pt(10.5), bold=True)
        _set_cell_font(table.rows[i].cells[1], size=Pt(10.5))
        _set_cell_shading(table.rows[i].cells[0], "F0F3F8")
    _set_table_col_widths(table, [Cm(4.0), Cm(11.6)])


def _set_table_header(header_row, labels: list[str]):
    """设置表格表头行样式。"""
    for i, label in enumerate(labels):
        header_row.cells[i].text = label
        _set_cell_font(header_row.cells[i], size=Pt(9.5), bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        _set_cell_shading(header_row.cells[i], "1F6FEB")


def _set_cell_font(cell, size=None, bold=False, color=None):
    """设置单元格字体。"""
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_before = Pt(2)
        paragraph.paragraph_format.space_after = Pt(2)
        for run in paragraph.runs:
            run.font.name = "微软雅黑"
            run.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
            if size:
                run.font.size = size
            if bold:
                run.bold = True
            if color:
                run.font.color.rgb = color


def _set_cell_shading(cell, color_hex: str):
    """设置单元格背景色。"""
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def _set_table_col_widths(table, widths: list):
    """设置表格列宽。"""
    for row in table.rows:
        for i, width in enumerate(widths):
            if i < len(row.cells):
                row.cells[i].width = width


def _add_horizontal_line(doc):
    """添加水平分隔线。"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CCCCCC")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _risk_cn(risk: str) -> str:
    return {"safe": "安全", "suspicious": "可疑", "malicious": "恶意"}.get(risk, risk)


def _risk_color_hex(risk: str) -> str:
    return {"safe": "E7F5EE", "suspicious": "FFF1CF", "malicious": "FFE0DF"}.get(risk, "F5F6F8")


def _format_size(size: int) -> str:
    if size < 1024:
        return f"{size} B"
    elif size < 1024 * 1024:
        return f"{size / 1024:.1f} KB"
    elif size < 1024 * 1024 * 1024:
        return f"{size / (1024 * 1024):.1f} MB"
    return f"{size / (1024 * 1024 * 1024):.2f} GB"


def _format_iso(iso_str: str) -> str:
    """将 ISO 时间字符串转为本地可读格式。"""
    try:
        dt = datetime.fromisoformat(iso_str)
        return f"{dt.year}年{dt.month:02d}月{dt.day:02d}日 {dt.hour:02d}:{dt.minute:02d}:{dt.second:02d}"
    except Exception:
        return iso_str


def _format_local_time() -> str:
    now = datetime.now()
    return f"{now.year}年{now.month:02d}月{now.day:02d}日 {now.hour:02d}:{now.minute:02d}:{now.second:02d}"
